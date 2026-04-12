"""
Gerador de Relatório CNIS em DOCX (Word)
Gera documento editável com a mesma estrutura do PDF.

O DOCX mantém:
- Identidade visual (cores, fontes onde possível)
- Estrutura de 12 seções
- Tabelas formatadas
- Textos dinâmicos e estáticos

O DOCX simplifica:
- Decorativos complexos (aspas gigantes, faixas diagonais) → substituídos por bordas e linhas
- Monograma SVG → substituído por texto estilizado "TS"
- Cards escuros → substituídos por tabelas com fundo
"""

import json
import sys
import io
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================================
#  PALETA DE CORES
# ============================================================================

COR_PRINCIPAL = RGBColor(0x1A, 0x3C, 0x40)      # Teal escuro
COR_DESTAQUE = RGBColor(0xE8, 0xB8, 0x8A)        # Dourado/pêssego
COR_TEXTO_ESCURO = RGBColor(0x1A, 0x1A, 0x1A)    # Preto suave
COR_BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
COR_TEAL_CLARO = RGBColor(0x2A, 0x6B, 0x6B)
COR_CINZA_CLARO = RGBColor(0xE5, 0xE5, 0xE5)
COR_VERMELHO = RGBColor(0xB7, 0x1C, 0x1C)
COR_LARANJA = RGBColor(0xE6, 0x51, 0x00)
COR_VERDE = RGBColor(0x2E, 0x7D, 0x32)


# ============================================================================
#  UTILITÁRIOS DOCX
# ============================================================================

def set_cell_shading(cell, color_hex: str):
    """Define cor de fundo de uma célula de tabela."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    """Configura espaçamento de parágrafo."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_run(paragraph, text, bold=False, italic=False, size=11,
            color=None, font_name=None):
    """Adiciona um run formatado a um parágrafo."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    return run


def add_heading_styled(doc, text, level=1, color=None, alignment=None,
                       font_size=None, font_name=None):
    """Adiciona um heading com estilo customizado."""
    heading = doc.add_heading(text, level=level)
    if alignment:
        heading.alignment = alignment
    for run in heading.runs:
        if color:
            run.font.color.rgb = color
        if font_size:
            run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
    return heading


def add_horizontal_line(doc, color_hex="1A3C40", width=1):
    """Adiciona uma linha horizontal decorativa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{width * 8}" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_page_break(doc):
    """Adiciona quebra de página."""
    doc.add_page_break()


def create_table_no_borders(doc, rows, cols):
    """Cria tabela sem bordas visíveis."""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remover bordas
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    return table


# ============================================================================
#  SEÇÕES DO RELATÓRIO
# ============================================================================

def pagina_capa(doc):
    """Página 1 — Capa."""
    # Espaço superior
    for _ in range(6):
        doc.add_paragraph()

    # Monograma (texto estilizado)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'TS', bold=True, size=72, color=COR_DESTAQUE, font_name='Georgia')

    doc.add_paragraph()

    # Nome do escritório
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'TATIANA SAMPAIO', bold=True, size=28, color=COR_DESTAQUE,
            font_name='Georgia')

    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'ADVOCACIA E CONSULTORIA JURÍDICA', size=10, color=COR_DESTAQUE)

    for _ in range(3):
        doc.add_paragraph()

    # Faixa com título
    add_horizontal_line(doc, 'E8B88A', 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'ANÁLISE DE CNIS', bold=True, size=24, color=COR_PRINCIPAL)
    add_horizontal_line(doc, 'E8B88A', 3)

    add_page_break(doc)


def pagina_apresentacao(doc):
    """Página 2 — Apresentação."""
    add_heading_styled(doc, 'Apresentação', level=1, color=COR_PRINCIPAL,
                       font_size=24, font_name='Georgia')
    add_horizontal_line(doc, 'E8B88A', 2)

    textos = [
        'O escritório Tatiana Sampaio Advocacia iniciou em 2006 e hoje é um moderno escritório de advocacia situado no Shopping Mestre Álvaro, em Serra/ES, com atuação em todo o território nacional.',
        'Nossa equipe é acolhedora e preparada para auxiliar nossos Clientes a se organizarem diante das demandas administrativas e judiciais.',
        'Com ferramentas modernas disponíveis para uma atuação estratégica na área de Direito Previdenciário, buscamos sempre o sucesso do nosso Cliente.',
    ]

    for texto in textos:
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
        for run in p.runs:
            run.font.size = Pt(12)

    doc.add_paragraph()

    # Assinatura
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, 'Dra. Tatiana Sampaio', bold=True, size=13, color=COR_PRINCIPAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, 'OAB/ES 12.297', bold=True, size=11, color=COR_PRINCIPAL)

    doc.add_paragraph()
    add_horizontal_line(doc, '1A3C40', 1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '+2000 FAMÍLIAS TRANSFORMADAS!', bold=True, size=16, color=COR_PRINCIPAL)

    add_page_break(doc)


def pagina_carta(doc):
    """Página 3 — Carta ao Cliente."""
    add_heading_styled(doc, 'Prezado(a) Sr(a),', level=1, color=COR_TEXTO_ESCURO,
                       font_size=22, font_name='Georgia')

    paragrafos = [
        'Dentre as medidas mais importantes para garantir segurança no momento de solicitar um benefício no INSS, a análise de CNIS se destaca como etapa essencial.',
        'Esta análise apresentará uma visão detalhada do seu extrato previdenciário (CNIS), com o objetivo de identificar possíveis inconsistências, tais como vínculos não registrados, datas incorretas, salários divergentes ou contribuições que possam não estar sendo consideradas pelo INSS.',
        'A correção prévia dessas informações é fundamental para evitar indeferimentos, atrasos na concessão do benefício ou até mesmo a concessão com valor inferior ao devido.',
    ]

    for texto in paragrafos:
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_paragraph_spacing(p, before=8, after=8, line_spacing=1.5)
        for run in p.runs:
            run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, 'Agradecemos pela confiança e prestígio que nos foram depositados.',
            italic=True, size=12, color=COR_TEXTO_ESCURO)

    add_page_break(doc)


def pagina_dados_segurado(doc, dados):
    """Página 4 — Dados do Segurado (dinâmica)."""
    cabecalho = dados.get('cabecalho', {})
    idade = dados.get('idade', {})

    add_heading_styled(doc, 'Dados do Segurado(a)', level=1, color=COR_PRINCIPAL,
                       font_size=22, font_name='Georgia')
    add_horizontal_line(doc, 'E8B88A', 2)

    # Nome em destaque
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_spacing(p, before=12, after=12)
    add_run(p, cabecalho.get('nome', 'NOME NÃO IDENTIFICADO'),
            bold=True, size=18, color=COR_PRINCIPAL)

    doc.add_paragraph()

    # Tabela de dados
    campos = [
        ('Data de Nascimento', cabecalho.get('data_nascimento', 'N/I')),
        ('CPF', cabecalho.get('cpf', 'N/I')),
        ('Idade', idade.get('descricao', 'N/I') if idade else 'N/I'),
        ('NIT', cabecalho.get('nit', 'N/I')),
    ]

    table = doc.add_table(rows=len(campos), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, valor) in enumerate(campos):
        # Célula do label
        cell_label = table.cell(i, 0)
        cell_label.width = Cm(6)
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, label, bold=True, size=12, color=COR_PRINCIPAL)
        set_cell_shading(cell_label, 'F5F5F5')

        # Célula do valor
        cell_valor = table.cell(i, 1)
        cell_valor.width = Cm(10)
        p = cell_valor.paragraphs[0]
        add_run(p, valor, size=13, color=COR_TEXTO_ESCURO)
        set_cell_shading(cell_valor, 'FFFFFF')

    add_page_break(doc)


def pagina_qualidade_segurado(doc, dados):
    """Página 5 — Qualidade de Segurado (dinâmica)."""
    qualidade = dados.get('qualidade_segurado', {})
    indicadores = dados.get('indicadores', {})
    data_analise = dados.get('data_analise', date.today().strftime('%d/%m/%Y'))

    add_heading_styled(doc, 'Qualidade de Segurado e Período de Graça', level=1,
                       color=COR_PRINCIPAL, font_size=22, font_name='Georgia')
    add_horizontal_line(doc, 'E8B88A', 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f'Com base no extrato, em {data_analise}:', bold=True, size=12,
            color=COR_TEXTO_ESCURO)

    doc.add_paragraph()

    # Status Atual
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, 'Status Atual: ', bold=True, size=12, color=COR_PRINCIPAL)
    add_run(p, qualidade.get('mensagem', 'Não determinado.'), size=12)

    # Período de Graça
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, 'Período de Graça: ', bold=True, size=12, color=COR_PRINCIPAL)

    status = qualidade.get('status', '')
    periodo = qualidade.get('periodo_graca_meses', 0)
    data_perda = qualidade.get('data_perda_estimada', '')

    if status == 'MANTIDA':
        add_run(p, f'O segurado mantém os direitos previdenciários por {periodo} meses '
                f'após a última contribuição. O período de graça se estende até {data_perda}.',
                size=12)
    elif status == 'PERDIDA':
        add_run(p, f'O período de graça de {periodo} meses já expirou. '
                'Para recuperar a qualidade de segurado, são necessárias '
                '6 contribuições válidas consecutivas.', size=12)
    else:
        add_run(p, 'Não foi possível determinar com os dados disponíveis.', size=12)

    # Alerta se houver
    if indicadores.get('total_alertas', 0) > 0:
        doc.add_paragraph()
        p = doc.add_paragraph()
        add_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, 'Atenção: ', bold=True, size=12, color=COR_LARANJA)
        add_run(p, f'Foram identificados {indicadores["total_alertas"]} alerta(s) '
                'que podem impactar a qualidade de segurado em períodos específicos.',
                size=12)

    add_page_break(doc)


def pagina_vinculos(doc, dados):
    """Página 6 — Vínculos Identificados (dinâmica)."""
    vinculos = dados.get('vinculos', [])
    data_analise = dados.get('data_analise', date.today().strftime('%d/%m/%Y'))

    add_heading_styled(doc, 'VÍNCULOS IDENTIFICADOS', level=1, color=COR_PRINCIPAL,
                       font_size=20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f'Data do Cálculo: {data_analise}', italic=True, size=11,
            color=COR_PRINCIPAL)

    doc.add_paragraph()

    if not vinculos:
        p = doc.add_paragraph('Nenhum vínculo identificado no extrato.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_break(doc)
        return

    # Tabela de vínculos
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabeçalho
    headers = ['Nº', 'Empregador / Atividade', 'Início', 'Fim', 'Situação']
    widths = [Cm(1.2), Cm(6.5), Cm(2.8), Cm(2.8), Cm(3.2)]

    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = width
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header, bold=True, size=9, color=COR_BRANCO)
        set_cell_shading(cell, '1A3C40')

    # Dados
    for vinculo in vinculos:
        row = table.add_row()

        # Nº
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(vinculo.get('seq', '')), bold=True, size=10)

        # Empregador
        p = row.cells[1].paragraphs[0]
        if vinculo.get('eh_beneficio'):
            nb = vinculo.get('numero_beneficio', '')
            esp = vinculo.get('especie_beneficio', 'Benefício')
            add_run(p, f'NB {nb} — {esp}', size=10)
        else:
            add_run(p, vinculo.get('empregador', 'Não identificado'), size=10)
        if vinculo.get('identificador_empregador'):
            p2 = row.cells[1].add_paragraph()
            add_run(p2, vinculo['identificador_empregador'], size=8,
                    color=RGBColor(0x77, 0x77, 0x77))

        # Início
        p = row.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, vinculo.get('data_inicio', '—'), size=10)

        # Fim
        p = row.cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        data_fim = vinculo.get('data_fim', '')
        if data_fim:
            add_run(p, data_fim, size=10)
        else:
            add_run(p, 'Em aberto', size=10, color=COR_VERDE, bold=True)

        # Situação
        p = row.cells[4].paragraphs[0]
        if vinculo.get('eh_beneficio'):
            sit = vinculo.get('situacao_beneficio', 'N/I')
            cor = COR_VERDE if sit == 'ATIVO' else COR_VERMELHO
            add_run(p, sit, size=10, color=cor)
        else:
            add_run(p, vinculo.get('tipo', 'N/I'), size=10)

    add_page_break(doc)


def pagina_indicadores(doc, dados):
    """Página 7 — Análise de Indicadores (dinâmica)."""
    indicadores = dados.get('indicadores', {})

    add_heading_styled(doc, 'Análise de Indicadores', level=1, color=COR_PRINCIPAL,
                       font_size=20, font_name='Georgia')
    add_horizontal_line(doc, 'E8B88A', 2)

    # A) Pendências
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=12, after=6)
    add_run(p, 'A) Grupo de Pendências (Necessitam de Prova ou Ajuste)',
            bold=True, size=13, color=COR_TEXTO_ESCURO)

    pendencias = indicadores.get('pendencias', [])
    if pendencias:
        for pend in pendencias:
            p = doc.add_paragraph()
            add_paragraph_spacing(p, before=4, after=4, line_spacing=1.4)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_run(p, pend['codigo'], bold=True, size=11, color=COR_PRINCIPAL)
            add_run(p, f' ({pend["nome"]}): ', italic=True, size=11,
                    color=RGBColor(0x55, 0x55, 0x55))
            add_run(p, pend['descricao'], size=11)

            # Significado
            p2 = doc.add_paragraph()
            add_paragraph_spacing(p2, before=0, after=6, line_spacing=1.3)
            add_run(p2, '    Significado: ', bold=True, size=10,
                    color=RGBColor(0x44, 0x44, 0x44))
            add_run(p2, pend['impacto'], size=10, color=RGBColor(0x44, 0x44, 0x44))
    else:
        p = doc.add_paragraph()
        add_run(p, 'Nenhuma pendência formal (P) foi identificada nas colunas de '
                'indicadores do extrato.', italic=True, size=11,
                color=RGBColor(0x88, 0x88, 0x88))

        p = doc.add_paragraph()
        add_paragraph_spacing(p, before=6, after=6)
        add_run(p, 'Nota: ', bold=True, size=10, color=COR_LARANJA)
        add_run(p, 'A ausência de indicadores no extrato resumido não garante que o '
                'vínculo esteja 100% saneado — recomenda-se verificar o Extrato '
                'Analítico no Meu INSS.', size=10)

    # B) Alertas
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=12, after=6)
    add_run(p, 'B) Grupo de Alertas (Informativos sobre o Regime)',
            bold=True, size=13, color=COR_TEXTO_ESCURO)

    alertas = indicadores.get('alertas', [])
    if alertas:
        for alerta in alertas:
            p = doc.add_paragraph()
            add_paragraph_spacing(p, before=4, after=4, line_spacing=1.4)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_run(p, alerta['codigo'], bold=True, size=11, color=COR_PRINCIPAL)
            add_run(p, f' ({alerta["nome"]}): ', italic=True, size=11,
                    color=RGBColor(0x55, 0x55, 0x55))
            add_run(p, alerta['descricao'], size=11)

            p2 = doc.add_paragraph()
            add_paragraph_spacing(p2, before=0, after=6, line_spacing=1.3)
            add_run(p2, '    Impacto: ', bold=True, size=10,
                    color=RGBColor(0x44, 0x44, 0x44))
            add_run(p2, alerta['impacto'], size=10, color=RGBColor(0x44, 0x44, 0x44))
    else:
        p = doc.add_paragraph()
        add_run(p, 'Nenhum alerta (I) foi identificado.', italic=True, size=11,
                color=RGBColor(0x88, 0x88, 0x88))

    # C) Acertos
    acertos = indicadores.get('acertos', [])
    if acertos:
        p = doc.add_paragraph()
        add_paragraph_spacing(p, before=12, after=6)
        add_run(p, 'C) Grupo de Acertos (Validações Concluídas)',
                bold=True, size=13, color=COR_TEXTO_ESCURO)
        for acerto in acertos:
            p = doc.add_paragraph()
            add_paragraph_spacing(p, before=4, after=4, line_spacing=1.4)
            add_run(p, acerto['codigo'], bold=True, size=11, color=COR_PRINCIPAL)
            add_run(p, f' ({acerto["nome"]}): ', italic=True, size=11)
            add_run(p, acerto['descricao'], size=11)

    # Desconhecidos
    desconhecidos = indicadores.get('desconhecidos', [])
    if desconhecidos:
        p = doc.add_paragraph()
        add_paragraph_spacing(p, before=12, after=6)
        add_run(p, 'Indicadores Não Identificados', bold=True, size=13,
                color=COR_LARANJA)

        for desc in desconhecidos:
            p = doc.add_paragraph()
            add_paragraph_spacing(p, before=4, after=4, line_spacing=1.4)
            add_run(p, f'{desc["codigo"]} ', bold=True, size=11, color=COR_LARANJA)
            add_run(p, '[DESCONHECIDO] ', bold=True, size=9, color=COR_LARANJA)
            add_run(p, f'— {desc["descricao"]}', size=11)

    add_page_break(doc)


def pagina_remuneracoes(doc, dados):
    """Página 8 — Remunerações Análise Crítica (dinâmica)."""
    remuneracoes = dados.get('remuneracoes', {})

    add_heading_styled(doc, 'REMUNERAÇÕES — ANÁLISE CRÍTICA', level=1,
                       color=COR_PRINCIPAL, font_size=20)
    add_horizontal_line(doc, 'E8B88A', 2)
    doc.add_paragraph()

    abaixo = remuneracoes.get('abaixo_minimo', [])
    proporcionais = remuneracoes.get('proporcionais', [])
    moeda_antiga = remuneracoes.get('moeda_antiga', [])

    if abaixo or proporcionais:
        # Tabela
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['Competência', 'Remuneração', 'Observação']
        widths = [Cm(3), Cm(3.5), Cm(10)]

        for i, (header, width) in enumerate(zip(headers, widths)):
            cell = table.rows[0].cells[i]
            cell.width = width
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, header, bold=True, size=9, color=COR_BRANCO)
            set_cell_shading(cell, '1A3C40')

        for rem in abaixo:
            row = table.add_row()
            p = row.cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, rem['competencia'], size=10)

            p = row.cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, f"R$ {rem['valor']:.2f}", size=10)

            p = row.cells[2].paragraphs[0]
            add_run(p, f"Abaixo do salário mínimo (SM = R$ {rem['salario_minimo']:.2f})",
                    size=10, bold=True, color=COR_VERMELHO)

        for rem in proporcionais:
            row = table.add_row()
            p = row.cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, rem['competencia'], size=10)

            p = row.cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, f"R$ {rem['valor']:.2f}", size=10)

            p = row.cells[2].paragraphs[0]
            add_run(p, rem.get('nota', 'Proporcional'), size=10, color=COR_LARANJA)

    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, 'Todas as remunerações analisadas estão dentro do salário '
                'mínimo vigente ou acima dele.', size=12, color=COR_VERDE)

    doc.add_paragraph()

    # Análise detalhada
    for rem in abaixo:
        p = doc.add_paragraph()
        add_paragraph_spacing(p, before=4, after=4, line_spacing=1.4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f"Competência {rem['competencia']} (R$ {rem['valor']:.2f}): ",
                bold=True, size=11, color=COR_PRINCIPAL)
        add_run(p, f"Esta remuneração está a {rem.get('percentual', 0)}% do salário "
                f"mínimo vigente (R$ {rem['salario_minimo']:.2f}). Sem agrupamento ou "
                "complementação, esta competência não contará como competência "
                "válida para carência.", size=11)

    for rem in proporcionais:
        p = doc.add_paragraph()
        add_paragraph_spacing(p, before=4, after=4, line_spacing=1.4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f"Competência {rem['competencia']} (R$ {rem['valor']:.2f}): ",
                bold=True, size=11, color=COR_PRINCIPAL)
        add_run(p, f"{rem.get('nota', '')}. Por se tratar de remuneração proporcional, "
                "não há impacto grave, mas recomenda-se o agrupamento de competências "
                "quando possível.", size=11)

    if moeda_antiga:
        doc.add_paragraph()
        p = doc.add_paragraph()
        add_paragraph_spacing(p, before=4, after=4, line_spacing=1.4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, 'Remunerações em moeda antiga: ', bold=True, size=11,
                color=COR_PRINCIPAL)
        add_run(p, 'As remunerações do período anterior a 07/1994 estão registradas '
                'em moeda antiga. Recomenda-se verificar a correta conversão para o '
                'Real junto ao INSS.', size=11)

    add_page_break(doc)


def pagina_beneficios(doc):
    """Página 9 — Benefícios Previdenciários (estática)."""
    add_heading_styled(doc, 'BENEFÍCIOS PREVIDENCIÁRIOS', level=1,
                       color=COR_PRINCIPAL, font_size=20)
    add_horizontal_line(doc, '1A3C40', 2)
    doc.add_paragraph()

    p = doc.add_paragraph(
        'Principais benefícios garantidos pela Seguridade Social brasileira ao '
        'segurado, caso mantenha a qualidade de segurado:'
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_paragraph_spacing(p, before=6, after=12, line_spacing=1.5)
    for run in p.runs:
        run.font.size = Pt(12)

    beneficios = [
        'Aposentadoria programada;',
        'Benefícios por incapacidade temporária ou permanente (auxílio-doença ou aposentadoria por invalidez);',
        'Auxílio Reclusão;',
        'Salário-Maternidade;',
        'Pensão por morte aos dependentes.',
    ]

    for b in beneficios:
        p = doc.add_paragraph(b, style='List Bullet')
        add_paragraph_spacing(p, before=3, after=3, line_spacing=1.4)
        for run in p.runs:
            run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph(
        'Um exemplo prático para compreender a importância da manutenção da '
        'qualidade de segurado, é do seguro de carro. Primeiro pagamos o seguro '
        'em dia e após isso, estamos assegurados em caso de algum sinistro. Não '
        'é à toa que o INSS é o Instituto Nacional do Seguro Social. Realizando '
        'a manutenção da qualidade de segurado, você garante o seu futuro em '
        'caso de doenças ou acidentes e garante uma qualidade de vida para seus '
        'familiares em caso de óbito.'
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_paragraph_spacing(p, before=6, after=6, line_spacing=1.5)
    for run in p.runs:
        run.font.size = Pt(12)

    add_page_break(doc)


def pagina_conclusao(doc, dados):
    """Página 10 — Conclusão (dinâmica)."""
    cabecalho = dados.get('cabecalho', {})
    qualidade = dados.get('qualidade_segurado', {})
    indicadores = dados.get('indicadores', {})
    lacunas = dados.get('lacunas', {})

    add_heading_styled(doc, 'CONCLUSÃO', level=1, color=COR_PRINCIPAL, font_size=20)
    add_horizontal_line(doc, '1A3C40', 2)
    doc.add_paragraph()

    nome = cabecalho.get('nome', 'N/I')
    status = qualidade.get('status', '')
    data_perda = qualidade.get('data_perda_estimada', '')

    # § 1 — Status
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=4, after=8, line_spacing=1.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, f'Diante da análise realizada, conclui-se que o(a) Sr(a). ', size=12)
    add_run(p, nome, bold=True, size=12, color=COR_PRINCIPAL)
    if status == 'MANTIDA':
        add_run(p, ' possui qualidade de segurado(a)', bold=True, size=12)
        add_run(p, f', com período de graça até {data_perda}.', size=12)
    elif status == 'PERDIDA':
        add_run(p, ' perdeu a qualidade de segurado(a)', bold=True, size=12,
                color=COR_VERMELHO)
        add_run(p, f' (desde aproximadamente {data_perda}).', size=12)
    else:
        add_run(p, ' tem situação indeterminada quanto à qualidade de segurado.',
                size=12)

    # § 2 — Recuperação (se perdeu)
    if status == 'PERDIDA':
        p = doc.add_paragraph()
        add_paragraph_spacing(p, before=4, after=8, line_spacing=1.6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, 'Para recuperar a qualidade de segurado(a), precisará de ',
                size=12)
        add_run(p, '6 contribuições válidas consecutivas', bold=True, size=12)
        add_run(p, ' (50% da carência de 12 meses para auxílio por incapacidade).',
                size=12)

    # § 3 — Pendências
    total_pend = indicadores.get('total_pendencias', 0)
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=4, after=8, line_spacing=1.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if total_pend > 0:
        add_run(p, f'Verificou-se ', size=12)
        add_run(p, f'{total_pend} pendência(s)', bold=True, size=12)
        add_run(p, ' que podem impactar o valor e a data de concessão do '
                'benefício no futuro. Recomenda-se a regularização junto ao '
                'INSS com a documentação apropriada.', size=12)
    else:
        add_run(p, 'O extrato não apresentou indicadores de pendência formal nas '
                'colunas de indicadores. Recomenda-se, contudo, a verificação do '
                'Extrato Analítico completo pelo portal Meu INSS.', size=12)

    # § 4 — Lacunas
    total_lacunas = lacunas.get('total', 0)
    if total_lacunas > 0:
        p = doc.add_paragraph()
        add_paragraph_spacing(p, before=4, after=8, line_spacing=1.6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'Foram identificadas ', size=12)
        add_run(p, f'{total_lacunas} lacuna(s) contributiva(s)', bold=True, size=12)
        maior = lacunas.get('maior_lacuna_meses', 0)
        add_run(p, f', sendo a maior de {maior} meses, que pode impactar a '
                'contagem de carência.', size=12)

    # § 5 — Disclaimer
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=4, after=8, line_spacing=1.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, 'Ressaltamos que todas as conclusões aqui apresentadas estão '
            'fundamentadas na legislação previdenciária vigente, podendo sofrer '
            'alterações caso o INSS ou o legislador promovam mudanças nas regras '
            'atualmente aplicáveis.', size=12)

    # § 6 — Fonte
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=4, after=8, line_spacing=1.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, 'Informo, ainda, que a análise foi realizada a partir da '
            'documentação apresentada pelo requerente, qual seja, CNIS.', size=12)

    # § 7 — Disponibilidade
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=4, after=8, line_spacing=1.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, 'Colocamo-nos à disposição para auxiliar em todas as etapas do '
            'processo e esclarecer quaisquer dúvidas que possam surgir durante a '
            'correção das pendências analisadas nesta análise de CNIS, estando '
            'nossa equipe disponível, especialmente, nas próximas 48 horas úteis '
            'para sanar todas as dúvidas.', size=12)

    doc.add_paragraph()
    add_horizontal_line(doc, 'E8B88A', 1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'Em caso de dúvidas estamos à disposição!', italic=True,
            size=13, color=COR_PRINCIPAL)

    add_page_break(doc)


def pagina_atuacao(doc):
    """Página 11 — Área de Atuação (estática)."""
    add_heading_styled(doc, 'NOSSA ÁREA DE ATUAÇÃO', level=1, color=COR_PRINCIPAL,
                       font_size=20)
    add_horizontal_line(doc, '1A3C40', 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_spacing(p, before=12, after=6)
    add_run(p, 'A área de atuação do Escritório Tatiana Sampaio Advocacia é: ',
            size=12)
    add_run(p, 'Direito Previdenciário', bold=True, size=14, color=COR_DESTAQUE)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'DIREITO PREVIDENCIÁRIO', bold=True, size=18, color=COR_PRINCIPAL)

    p = doc.add_paragraph(
        'Atuamos de forma especializada em todas as esferas do Direito '
        'Previdenciário, buscando garantir os direitos dos nossos clientes '
        'junto ao INSS, seja na via administrativa ou judicial.'
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_spacing(p, before=6, after=12, line_spacing=1.5)
    for run in p.runs:
        run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'BENEFÍCIOS DO INSS', bold=True, size=14, color=COR_PRINCIPAL)

    doc.add_paragraph()

    beneficios = [
        'Aposentadoria por Idade', 'Aposentadoria Especial',
        'Aposentadoria por Invalidez', 'Auxílio Acidente',
        'Pensão por Morte', 'Auxílio Doença',
        'BPC/LOAS Deficiente', 'Aposentadoria por Tempo de Contribuição',
        'Salário Maternidade', 'BPC/LOAS Idoso',
    ]

    # Tabela 2 colunas para os benefícios
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, beneficio in enumerate(beneficios):
        row_idx = i // 2
        col_idx = i % 2
        cell = table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, beneficio, size=11, color=COR_PRINCIPAL)
        set_cell_shading(cell, 'E8F5E9')

    add_page_break(doc)


def pagina_contracapa(doc):
    """Página 12 — Contracapa (estática)."""
    for _ in range(5):
        doc.add_paragraph()

    # Monograma
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'TS', bold=True, size=60, color=COR_DESTAQUE, font_name='Georgia')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'TATIANA SAMPAIO', bold=True, size=24, color=COR_DESTAQUE,
            font_name='Georgia')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'ADVOCACIA E CONSULTORIA JURÍDICA', size=10, color=COR_DESTAQUE)

    doc.add_paragraph()
    doc.add_paragraph()

    add_horizontal_line(doc, '2A6B6B', 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'OBRIGADA PELA CONFIANÇA', bold=True, size=18, color=COR_PRINCIPAL)
    add_horizontal_line(doc, '2A6B6B', 3)

    doc.add_paragraph()

    contato = [
        'Para mais informações:',
        'Tatiana Sampaio Advocacia',
        'Whatsapp: 27 99694-5544',
        'Instagram: @dratatianasampaio',
        'TIKTOK: @dra.aposentadoria',
        'Youtube: dratatianasampaio',
    ]

    for linha in contato:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bold = linha == contato[0] or linha == contato[1]
        add_run(p, linha, bold=bold, size=12, color=COR_TEXTO_ESCURO)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'Aqui, respeitamos a história do nosso Cliente!', italic=True,
            size=13, color=COR_DESTAQUE)


# ============================================================================
#  FUNÇÃO PRINCIPAL
# ============================================================================

def gerar_docx(dados_analise: dict) -> bytes:
    """Gera relatório DOCX completo.

    Args:
        dados_analise: Dicionário retornado pelo cnis_analyzer

    Returns:
        Bytes do arquivo DOCX
    """
    doc = Document()

    # Configurar margens A4
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Gerar cada página
    pagina_capa(doc)
    pagina_apresentacao(doc)
    pagina_carta(doc)
    pagina_dados_segurado(doc, dados_analise)
    pagina_qualidade_segurado(doc, dados_analise)
    pagina_vinculos(doc, dados_analise)
    pagina_indicadores(doc, dados_analise)
    pagina_remuneracoes(doc, dados_analise)
    pagina_beneficios(doc)
    pagina_conclusao(doc, dados_analise)
    pagina_atuacao(doc)
    pagina_contracapa(doc)

    # Salvar em bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def gerar_nome_arquivo_docx(dados_analise: dict) -> str:
    """Gera nome do arquivo DOCX."""
    nome = dados_analise.get('cabecalho', {}).get('nome', 'SEGURADO')
    data = date.today().strftime('%d-%m-%Y')
    nome_limpo = ''.join(c for c in nome if c.isalnum() or c in ' .-_')
    return f"Análise CNIS - {nome_limpo.strip()} - {data}.docx"


# ============================================================================
#  PONTO DE ENTRADA
# ============================================================================

if __name__ == '__main__':
    dados_stdin = sys.stdin.read()
    try:
        dados = json.loads(dados_stdin)
    except json.JSONDecodeError as e:
        print(f'JSON inválido: {e}', file=sys.stderr)
        sys.exit(1)

    docx_bytes = gerar_docx(dados)
    nome = gerar_nome_arquivo_docx(dados)

    with open(nome, 'wb') as f:
        f.write(docx_bytes)
    print(f'DOCX gerado: {nome}')
